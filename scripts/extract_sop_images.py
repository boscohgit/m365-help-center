#!/usr/bin/env python3
"""Extract body screenshots from SOP DOCX files and build review contact sheets."""

from __future__ import annotations

import hashlib
import json
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = Path(__file__).resolve().parent / ".image-review"
DOC_TO_SLUG = {
    "1-1.安装M365应用套件.docx": "install-m365",
    "2.Office365登录.docx": "office-sign-in",
    "3.Outlook邮箱登录.docx": "outlook",
    "5. Outlook使用指引.docx": "outlook",
    "6.Teams使用指引.docx": "teams",
    "7.OneDrive使用指引.docx": "onedrive",
    "Autopilot_新电脑首次登录及公司应用安装.docx": "autopilot-new-pc",
    "NAS服务器文件夹（扫描Scan文件夹）访问配置.docx": "scan-documents",
    "使用 SharePoint 共享文件指引.docx": "sharepoint",
    "分所账号配置.docx": "branch-account",
    "自主重置公司账号密码.docx": "password-reset",
    "自助修改密码.docx": "password-reset",
    "4.添加账号验证方式及修改密码（优化版）.docx": "security-info",
}


def safe_name(value: str) -> str:
    return re.sub(r"[^a-zA-Z0-9._-]+", "-", value).strip("-").lower()


def nearby_text(paragraphs: list, index: int) -> str:
    for cursor in range(index, max(-1, index - 4), -1):
        text = paragraphs[cursor].text.strip()
        if text:
            return text[:160]
    return ""


def review_key(doc_path: Path, slug: str) -> str:
    suffix = hashlib.sha1(doc_path.name.encode("utf-8")).hexdigest()[:7]
    return f"{slug}-{suffix}"


def extract_document(doc_path: Path, slug: str) -> list[dict]:
    document = Document(doc_path)
    paragraphs = list(document.paragraphs)
    records: list[dict] = []
    seen_rel_ids: set[str] = set()
    doc_dir = OUTPUT / review_key(doc_path, slug)
    doc_dir.mkdir(parents=True, exist_ok=True)

    # document.paragraphs intentionally excludes headers and footers.
    for paragraph_index, paragraph in enumerate(paragraphs):
        for blip in paragraph._p.xpath(".//a:blip"):
            rel_id = blip.get(qn("r:embed"))
            if not rel_id or rel_id in seen_rel_ids:
                continue
            seen_rel_ids.add(rel_id)
            part = document.part.related_parts[rel_id]
            extension = Path(part.partname).suffix.lower() or ".png"
            number = len(records) + 1
            filename = f"{number:02d}{extension}"
            destination = doc_dir / filename
            destination.write_bytes(part.blob)
            digest = hashlib.sha256(part.blob).hexdigest()
            records.append(
                {
                    "document": doc_path.name,
                    "slug": slug,
                    "sequence": number,
                    "paragraph_index": paragraph_index,
                    "context": nearby_text(paragraphs, paragraph_index),
                    "source_media": Path(part.partname).name,
                    "review_path": str(destination),
                    "sha256": digest,
                }
            )
    return records


def make_contact_sheet(records: list[dict], destination: Path) -> None:
    if not records:
        return
    cell_width, cell_height = 360, 330
    columns = 3
    rows = (len(records) + columns - 1) // columns
    canvas = Image.new("RGB", (cell_width * columns, cell_height * rows), "white")
    draw = ImageDraw.Draw(canvas)
    font = ImageFont.load_default(size=18)
    small = ImageFont.load_default(size=13)
    for index, record in enumerate(records):
        image = Image.open(record["review_path"]).convert("RGB")
        image.thumbnail((cell_width - 20, cell_height - 82))
        x = (index % columns) * cell_width
        y = (index // columns) * cell_height
        image_x = x + (cell_width - image.width) // 2
        canvas.paste(image, (image_x, y + 34))
        draw.text((x + 10, y + 8), f'#{record["sequence"]:02d}', fill="black", font=font)
        context = record["context"].replace("\n", " ")[:46]
        draw.text((x + 10, y + cell_height - 36), context, fill="#333333", font=small)
    canvas.save(destination, quality=88)


def main() -> None:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    manifest: list[dict] = []
    for filename, slug in DOC_TO_SLUG.items():
        doc_path = ROOT / filename
        if not doc_path.exists():
            continue
        records = extract_document(doc_path, slug)
        manifest.extend(records)
        make_contact_sheet(records, OUTPUT / f"{review_key(doc_path, slug)}-contact.jpg")
        print(f"{filename}: {len(records)} body images")
    (OUTPUT / "manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    print(f"Total: {len(manifest)} body images")


if __name__ == "__main__":
    main()
